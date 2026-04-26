"""
docx_processor.py
=================
Extracts text from .docx files and returns it in the same block format
as PDFProcessor, so the rest of the pipeline (language detection,
translation, TTS, document analysis) works identically for both file types.

Block schema (mirrors PDFProcessor output):
    {
        "text":          str,    # stripped paragraph or table cell text
        "font_size":     float,  # 0.0 — DOCX doesn't map to PDF font sizes;
                                 # heading detection uses style name instead
        "page_num":      int,    # 0-based logical page (approximated)
        "page_height":   float,  # 792.0 (A4 default — unused for DOCX)
        "y0":            float,  # 0.0 (no geometry for DOCX)
        "y1":            float,  # 0.0
        "is_heading":    bool,   # True for Heading 1/2/3 styles
        "heading_level": int,    # 1/2/3 — 0 if not a heading
        "source":        str,    # "paragraph" | "table" | "header" | "footer"
    }

Heading detection uses Word's built-in Heading styles. Body text has
font_size=0.0 which causes split_into_chapters to use regex fallback —
but since we set is_heading=True directly, the font-size path is skipped
and heading-based splitting works correctly.

Table content is included as plain text (one block per row, cells joined
with " | "). Tables are placed after their surrounding paragraphs.

Headers and footers are extracted but tagged as source="header"/"footer"
so filter_blocks can optionally remove them (same as PDF margin zones).
"""

import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


# Heading style name → heading level (1 = top)
_HEADING_LEVEL_MAP: dict[str, int] = {
    "Title":      1,
    "Subtitle":   2,
    "Heading 1":  1,
    "Heading 2":  2,
    "Heading 3":  3,
    "Heading 4":  3,
    "Heading 5":  3,
    "Heading 6":  3,
}

# Approximate chars per "page" for logical page numbering
_CHARS_PER_PAGE = 3000


class DocxProcessor:
    """Extract and structure text from .docx files."""

    # ------------------------------------------------------------------ #
    # Public API                                                           #
    # ------------------------------------------------------------------ #

    @staticmethod
    def validate(file_bytes: bytes) -> tuple[bool, str]:
        """Return (ok, error_message). Checks magic bytes + parsability."""
        if not file_bytes:
            return False, "Empty file."
        # DOCX is a ZIP — starts with PK\x03\x04
        if not file_bytes[:4] == b"PK\x03\x04":
            return False, "Not a valid .docx file (bad magic bytes)."
        try:
            DocxProcessor._open(file_bytes)
            return True, ""
        except Exception as e:
            return False, f"Could not open .docx: {e}"

    @staticmethod
    def get_metadata(file_bytes: bytes) -> dict:
        """Extract core properties: title, author, subject, description."""
        try:
            doc   = DocxProcessor._open(file_bytes)
            props = doc.core_properties
            return {
                "title":       (props.title   or "").strip(),
                "author":      (props.author  or "").strip(),
                "subject":     (props.subject or "").strip(),
                "description": (props.description or "").strip(),
            }
        except Exception:
            return {}

    @staticmethod
    def get_page_count(file_bytes: bytes) -> int:
        """Approximate page count based on character count."""
        try:
            doc   = DocxProcessor._open(file_bytes)
            total = sum(len(p.text) for p in doc.paragraphs)
            return max(1, round(total / _CHARS_PER_PAGE))
        except Exception:
            return 1

    @staticmethod
    def extract(
        file_bytes:       bytes,
        lang_code:        str           = "en",
        include_tables:   bool          = True,
        include_hf:       bool          = False,   # headers/footers
    ) -> dict:
        """
        Extract all text from a .docx file.

        Returns:
            {
                "blocks":     list[dict],   # same schema as PDFProcessor
                "metadata":   dict,
                "page_count": int,
                "mode_used":  "docx",
            }
        """
        try:
            doc = DocxProcessor._open(file_bytes)
        except Exception as e:
            logger.error(f"DocxProcessor.extract failed to open file: {e}")
            return {"blocks": [], "metadata": {}, "page_count": 1, "mode_used": "docx"}

        metadata   = DocxProcessor.get_metadata(file_bytes)
        blocks     = []
        char_count = 0   # running char total for page estimation

        # Walk all body elements in document order
        # python-docx exposes doc.element.body which gives real order
        # but iterating doc.paragraphs + doc.tables separately loses order.
        # Use the XML body iterator instead.
        from docx.oxml.ns import qn
        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                # Paragraph
                para   = _para_from_element(doc, child)
                if para is None:
                    continue
                text   = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                level      = _HEADING_LEVEL_MAP.get(style_name, 0)
                page_num   = char_count // _CHARS_PER_PAGE

                blocks.append({
                    "text":          text,
                    "font_size":     0.0,
                    "page_num":      page_num,
                    "page_height":   792.0,
                    "y0":            0.0,
                    "y1":            0.0,
                    "is_heading":    level > 0,
                    "heading_level": level,
                    "source":        "paragraph",
                })
                char_count += len(text)

            elif tag == "tbl" and include_tables:
                # Table — one block per row, cells joined with " | "
                from docx.table import Table as DocxTable
                table = DocxTable(child, doc)
                for row in table.rows:
                    cells = [
                        c.text.strip()
                        for c in row.cells
                        if c.text.strip()
                    ]
                    # Deduplicate merged cells (python-docx repeats them)
                    seen, unique = set(), []
                    for cell in cells:
                        if cell not in seen:
                            seen.add(cell)
                            unique.append(cell)
                    row_text = " | ".join(unique)
                    if not row_text:
                        continue
                    page_num = char_count // _CHARS_PER_PAGE
                    blocks.append({
                        "text":          row_text,
                        "font_size":     0.0,
                        "page_num":      page_num,
                        "page_height":   792.0,
                        "y0":            0.0,
                        "y1":            0.0,
                        "is_heading":    False,
                        "heading_level": 0,
                        "source":        "table",
                    })
                    char_count += len(row_text)

        # Optional: headers and footers from all sections
        if include_hf:
            for section in doc.sections:
                for hf_para in list(section.header.paragraphs) + list(section.footer.paragraphs):
                    text = hf_para.text.strip()
                    if text:
                        src = "header" if hf_para in section.header.paragraphs else "footer"
                        blocks.append({
                            "text":          text,
                            "font_size":     0.0,
                            "page_num":      0,
                            "page_height":   792.0,
                            "y0":            0.0 if src == "header" else 750.0,
                            "y1":            0.0 if src == "header" else 792.0,
                            "is_heading":    False,
                            "heading_level": 0,
                            "source":        src,
                        })

        page_count = max(1, char_count // _CHARS_PER_PAGE)
        logger.info(
            f"DocxProcessor.extract: {len(blocks)} blocks, "
            f"~{page_count} pages, {char_count} chars"
        )

        return {
            "blocks":     blocks,
            "metadata":   metadata,
            "page_count": page_count,
            "mode_used":  "docx",
        }

    # ------------------------------------------------------------------ #
    # Helpers                                                              #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _open(file_bytes: bytes):
        """Open a .docx from bytes."""
        from docx import Document
        return Document(io.BytesIO(file_bytes))

    @staticmethod
    def is_docx(filename: str) -> bool:
        return filename.lower().endswith(".docx")

    @staticmethod
    def display_name_from_metadata(metadata: dict, filename: str) -> str:
        """Best display name: metadata title → cleaned filename."""
        title = (metadata.get("title") or "").strip()
        if title and len(title) > 4:
            return title[:80]
        name = filename.rsplit(".", 1)[0]
        name = re.sub(r"[_\-]+", " ", name).strip()
        return name[:80]


# ------------------------------------------------------------------ #
# Internal helpers                                                     #
# ------------------------------------------------------------------ #

def _para_from_element(doc, elem):
    """Wrap a raw XML paragraph element as a python-docx Paragraph."""
    try:
        from docx.text.paragraph import Paragraph
        return Paragraph(elem, doc)
    except Exception:
        return None